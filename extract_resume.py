#!/usr/bin/env python3
"""Script to extract text content from resume_.docx"""
try:
    from docx import Document
    import json
    
    doc = Document('resume_.docx')
    
    # Extract all paragraphs
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text.strip())
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                content.append(' | '.join(row_data))
    
    # Print extracted content
    print('\n'.join(content))
    
    # Save to JSON for easy use
    with open('resume_content.json', 'w', encoding='utf-8') as f:
        json.dump({'content': content}, f, indent=2, ensure_ascii=False)
        
except ImportError:
    print("python-docx not installed. Install it with: pip install python-docx")
except Exception as e:
    print(f"Error: {e}")

